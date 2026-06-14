import os
import io
import gc
import pypdfium2 as pdfium  # Replaced PyMuPDF with pypdfium2 for better stability
from docx import Document
from services.pdf_parser import extract_with_pdfplumber, is_extraction_good


def extract_from_docx(docx_bytes: bytes) -> dict:
    """Extract text from DOCX file."""
    try:
        doc = Document(io.BytesIO(docx_bytes))
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        text = '\n\n'.join(paragraphs)
        return {
            "text": text if text else "Unable to extract text from DOCX.",
            "page_count": 1,
            "parser_used": "python-docx"
        }
    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")


def extract_with_pypdfium2(pdf_bytes: bytes) -> tuple[str, int]:
    """Layer 1: pypdfium2 — Fast C++ engine, lightweight memory (~15MB)."""
    doc = pdfium.PdfDocument(pdf_bytes)
    page_count = len(doc)
    text_parts = []
    for page in doc:
        text_page = page.get_textpage()
        text_parts.append(text_page.get_text_bounded().strip())
    return "\n\n".join(text_parts), page_count


def extract_resume_text(pdf_bytes: bytes) -> dict:
    """
    2-layer orchestrator:
    Layer 1: pypdfium2 (Fast, lightweight C++ engine)
    Layer 2: pdfplumber (Heavy Python heap, fallback)
    Note: OCR fallback was removed to prevent OOM memory issues on the server.
    """
    # Layer 1 — pypdfium2
    text, page_count = extract_with_pypdfium2(pdf_bytes)
    if is_extraction_good(text):
        return {"text": text, "page_count": page_count, "parser_used": "pypdfium2"}

    # Layer 2 — pdfplumber (Fallback)
    try:
        text, page_count = extract_with_pdfplumber(pdf_bytes)
        if is_extraction_good(text):
            return {"text": text, "page_count": page_count, "parser_used": "pdfplumber"}
    finally:
        gc.collect()

    raise ValueError("We couldn't read the text in this file. Please upload a standard digital PDF or Word document instead of a scanned image.")
